"""Parser for DOCX (Microsoft Word) documents."""

from __future__ import annotations

from pathlib import Path
from typing import Any

from rag_agent.models.document import Document, DocumentPage
from rag_agent.parsers.base import BaseDocumentParser


class DocxDocumentParser(BaseDocumentParser):
    """Parser for .docx files using python-docx."""

    SUPPORTED_EXTENSIONS = {".docx"}

    async def parse(self, filepath: Path) -> Document:
        """Parse a DOCX file.

        Raises:
            ValueError: If the file is not a DOCX file.
        """
        if not self.is_extension_allowed(filepath):
            raise ValueError(
                f"Extension {filepath.suffix} not supported by DocxDocumentParser"
            )

        from docx import Document as DOCXDocument

        doc: Any = DOCXDocument(str(filepath))
        content = "\n".join(p.text for p in doc.paragraphs)

        page = DocumentPage(page_num=1, content=content)
        return Document(pages=[page], metadata=self.get_document_metadata(filepath))
